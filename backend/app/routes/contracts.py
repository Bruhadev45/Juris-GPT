"""
Contract Generation API Routes for JurisGPT
Indian Law Compliant Contract Generation with AI
"""

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import Dict, List, Optional, Any
from datetime import datetime, timezone
from uuid import UUID, uuid4
from openai import OpenAI
from app.config import settings
from app.database import supabase
import io

# Initialize OpenAI client with validation
def get_openai_client():
    """Get OpenAI client, validating API key"""
    api_key = settings.openai_api_key
    if not api_key or api_key == "sk-placeholder" or len(api_key) < 20:
        raise ValueError("OpenAI API key is not configured properly")
    return OpenAI(api_key=api_key)

# Lazy initialization - client will be created on first use
_openai_client = None

def get_client():
    global _openai_client
    if _openai_client is None:
        _openai_client = get_openai_client()
    return _openai_client

router = APIRouter()


# ============== Contract Types Database ==============

CONTRACT_TYPES = {
    "nda": {
        "name": "Non-Disclosure Agreement",
        "category": "Startup Essentials",
        "description": "Protect confidential information",
        "law_reference": "Indian Contract Act, 1872",
        "fields": [
            "party1_name", "party1_address", "party2_name", "party2_address",
            "effective_date", "term_years", "confidential_info_description", "jurisdiction_state"
        ],
        "field_labels": {
            "party1_name": "Disclosing Party Name",
            "party1_address": "Disclosing Party Address",
            "party2_name": "Receiving Party Name",
            "party2_address": "Receiving Party Address",
            "effective_date": "Effective Date",
            "term_years": "Term (Years)",
            "confidential_info_description": "Description of Confidential Information",
            "jurisdiction_state": "Jurisdiction State"
        },
        "field_types": {
            "effective_date": "date",
            "term_years": "number"
        }
    },
    "employment": {
        "name": "Employment Agreement",
        "category": "Employment & HR",
        "description": "Hire employees with Indian labour law compliance",
        "law_reference": "Indian Labour Laws, PF/ESI Act",
        "fields": [
            "employer_name", "employer_address", "employer_gstin", "employee_name",
            "employee_address", "employee_pan", "designation", "department", "salary",
            "joining_date", "probation_months", "notice_period_days", "work_location"
        ],
        "field_labels": {
            "employer_name": "Employer/Company Name",
            "employer_address": "Employer Address",
            "employer_gstin": "Employer GSTIN",
            "employee_name": "Employee Name",
            "employee_address": "Employee Address",
            "employee_pan": "Employee PAN",
            "designation": "Job Designation",
            "department": "Department",
            "salary": "Monthly Salary (INR)",
            "joining_date": "Joining Date",
            "probation_months": "Probation Period (Months)",
            "notice_period_days": "Notice Period (Days)",
            "work_location": "Work Location"
        },
        "field_types": {
            "salary": "number",
            "joining_date": "date",
            "probation_months": "number",
            "notice_period_days": "number"
        }
    },
    "service_agreement": {
        "name": "Service Agreement (MSA)",
        "category": "Business Agreements",
        "description": "B2B service contracts with GST compliance",
        "law_reference": "Indian Contract Act, GST Act",
        "fields": [
            "provider_name", "provider_address", "provider_gstin", "client_name",
            "client_address", "client_gstin", "services_description", "payment_terms",
            "contract_value", "start_date", "end_date"
        ],
        "field_labels": {
            "provider_name": "Service Provider Name",
            "provider_address": "Provider Address",
            "provider_gstin": "Provider GSTIN",
            "client_name": "Client Name",
            "client_address": "Client Address",
            "client_gstin": "Client GSTIN",
            "services_description": "Description of Services",
            "payment_terms": "Payment Terms",
            "contract_value": "Contract Value (INR)",
            "start_date": "Start Date",
            "end_date": "End Date"
        },
        "field_types": {
            "contract_value": "number",
            "start_date": "date",
            "end_date": "date"
        }
    },
    "privacy_policy": {
        "name": "Privacy Policy",
        "category": "Compliance Documents",
        "description": "DPDP Act 2023 compliant privacy policy",
        "law_reference": "Digital Personal Data Protection Act, 2023",
        "fields": [
            "company_name", "website_url", "contact_email", "data_types_collected",
            "data_processing_purposes", "data_retention_period", "grievance_officer_name",
            "grievance_officer_email"
        ],
        "field_labels": {
            "company_name": "Company Name",
            "website_url": "Website URL",
            "contact_email": "Contact Email",
            "data_types_collected": "Types of Data Collected",
            "data_processing_purposes": "Data Processing Purposes",
            "data_retention_period": "Data Retention Period",
            "grievance_officer_name": "Grievance Officer Name",
            "grievance_officer_email": "Grievance Officer Email"
        },
        "field_types": {}
    },
    "terms_of_service": {
        "name": "Terms of Service",
        "category": "Compliance Documents",
        "description": "Website/app terms compliant with IT Act",
        "law_reference": "IT Act 2000, Consumer Protection Act",
        "fields": [
            "company_name", "website_url", "service_description", "user_obligations",
            "payment_terms", "refund_policy", "jurisdiction_state"
        ],
        "field_labels": {
            "company_name": "Company Name",
            "website_url": "Website URL",
            "service_description": "Service Description",
            "user_obligations": "User Obligations",
            "payment_terms": "Payment Terms",
            "refund_policy": "Refund Policy",
            "jurisdiction_state": "Jurisdiction State"
        },
        "field_types": {}
    },
    "founders_agreement": {
        "name": "Founder's Agreement",
        "category": "Startup Essentials",
        "description": "Agreement between co-founders",
        "law_reference": "Companies Act, 2013",
        "fields": [
            "company_name", "founder1_name", "founder1_equity", "founder2_name",
            "founder2_equity", "vesting_period", "cliff_months", "roles_responsibilities",
            "ip_assignment", "non_compete_years"
        ],
        "field_labels": {
            "company_name": "Company Name",
            "founder1_name": "Founder 1 Name",
            "founder1_equity": "Founder 1 Equity (%)",
            "founder2_name": "Founder 2 Name",
            "founder2_equity": "Founder 2 Equity (%)",
            "vesting_period": "Vesting Period (Months)",
            "cliff_months": "Cliff Period (Months)",
            "roles_responsibilities": "Roles & Responsibilities",
            "ip_assignment": "IP Assignment Clause",
            "non_compete_years": "Non-Compete Period (Years)"
        },
        "field_types": {
            "founder1_equity": "number",
            "founder2_equity": "number",
            "vesting_period": "number",
            "cliff_months": "number",
            "non_compete_years": "number"
        }
    },
    "consultant_agreement": {
        "name": "Consultant Agreement",
        "category": "Employment & HR",
        "description": "Hire consultants with TDS compliance",
        "law_reference": "Income Tax Act Section 194J",
        "fields": [
            "company_name", "company_address", "company_gstin", "consultant_name",
            "consultant_address", "consultant_pan", "services_scope", "fees",
            "payment_schedule", "start_date", "end_date", "tds_rate"
        ],
        "field_labels": {
            "company_name": "Company Name",
            "company_address": "Company Address",
            "company_gstin": "Company GSTIN",
            "consultant_name": "Consultant Name",
            "consultant_address": "Consultant Address",
            "consultant_pan": "Consultant PAN",
            "services_scope": "Scope of Services",
            "fees": "Consulting Fees (INR)",
            "payment_schedule": "Payment Schedule",
            "start_date": "Start Date",
            "end_date": "End Date",
            "tds_rate": "TDS Rate (%)"
        },
        "field_types": {
            "fees": "number",
            "start_date": "date",
            "end_date": "date",
            "tds_rate": "number"
        }
    },
    "board_resolution": {
        "name": "Board Resolution",
        "category": "Compliance Documents",
        "description": "Standard board resolution format",
        "law_reference": "Companies Act, 2013",
        "fields": [
            "company_name", "cin", "registered_office", "meeting_date",
            "resolution_subject", "resolution_text", "directors_present"
        ],
        "field_labels": {
            "company_name": "Company Name",
            "cin": "Company Identification Number (CIN)",
            "registered_office": "Registered Office Address",
            "meeting_date": "Meeting Date",
            "resolution_subject": "Resolution Subject",
            "resolution_text": "Resolution Text",
            "directors_present": "Directors Present"
        },
        "field_types": {
            "meeting_date": "date"
        }
    }
}

# Group contract types by category
def get_contracts_by_category():
    categories = {}
    for type_id, contract in CONTRACT_TYPES.items():
        category = contract["category"]
        if category not in categories:
            categories[category] = []
        categories[category].append({
            "id": type_id,
            "name": contract["name"],
            "description": contract["description"],
            "law_reference": contract["law_reference"]
        })
    return categories


# ============== Pydantic Models ==============

class ContractTypeResponse(BaseModel):
    id: str
    name: str
    category: str
    description: str
    law_reference: str
    fields: List[str]
    field_labels: Dict[str, str]
    field_types: Dict[str, str]


class ContractTypesListResponse(BaseModel):
    categories: Dict[str, List[Dict[str, str]]]
    total_count: int


class ContractGenerateRequest(BaseModel):
    contract_type: str = Field(..., description="Contract type ID (e.g., 'nda', 'employment')")
    form_data: Dict[str, Any] = Field(..., description="Form data with field values")
    language: Optional[str] = Field(default="English", description="Language for the contract")

    class Config:
        json_schema_extra = {
            "example": {
                "contract_type": "nda",
                "form_data": {
                    "disclosingPartyName": "ABC Pvt Ltd",
                    "disclosingPartyAddress": "123 MG Road, Bangalore 560001",
                    "receivingPartyName": "XYZ Technologies",
                    "receivingPartyAddress": "456 HSR Layout, Bangalore 560102",
                    "effectiveDate": "2024-01-15",
                    "confidentialityPeriod": 2,
                    "purpose": "Evaluation of potential business partnership",
                    "confidentialInfo": "Business plans, technical specifications, and customer data",
                    "governingLaw": "Karnataka"
                }
            }
        }


class GeneratedContractResponse(BaseModel):
    id: str
    contract_type: str
    contract_name: str
    content: str
    form_data: Dict[str, Any]
    created_at: str
    law_reference: str
    status: str = "generated"


class ContractListItem(BaseModel):
    id: str
    contract_type: str
    contract_name: str
    created_at: str
    status: str


# ============== In-Memory Storage (for demo - replace with Supabase in production) ==============

generated_contracts_store: Dict[str, Dict] = {}


# ============== AI Contract Generation ==============

def camel_to_title(name: str) -> str:
    """Convert camelCase field name to Title Case label"""
    import re
    # Insert space before uppercase letters and capitalize first letter
    result = re.sub(r'([A-Z])', r' \1', name)
    return result.strip().title()


def build_contract_prompt(contract_type: str, form_data: Dict[str, Any], language: str = "English") -> str:
    """Build the AI prompt for contract generation"""

    contract_info = CONTRACT_TYPES.get(contract_type)
    if not contract_info:
        raise ValueError(f"Unknown contract type: {contract_type}")

    # Format form data for the prompt - handle both backend field names and frontend camelCase names
    form_data_text = "\n".join([
        f"- {contract_info['field_labels'].get(key, camel_to_title(key))}: {value}"
        for key, value in form_data.items()
        if value is not None and value != "" and value is not False
    ])

    system_prompt = f"""You are an expert Indian corporate lawyer with 20+ years of experience drafting legal documents.
You specialize in Indian law and ensure all contracts comply with applicable Indian legislation.

Your task is to generate a professional, legally binding {contract_info['name']} that:
1. Complies with {contract_info['law_reference']}
2. Uses precise legal terminology appropriate for Indian jurisdiction
3. Includes all standard protective clauses
4. Is comprehensive yet clear and readable
5. Can be used directly after legal review

Important Indian Law Considerations:
- All monetary values should be in Indian Rupees (INR/Rs.)
- Dates should follow DD-MM-YYYY format
- Include appropriate stamp duty notice
- Reference applicable Indian laws and sections
- Include jurisdiction and governing law clauses
- Add dispute resolution mechanism (arbitration/mediation)
- Ensure compliance with latest amendments

Output Format:
- Use clear section headings with numbers (1. 2. 3. etc.)
- Include proper definitions section
- Add signature blocks for all parties
- Include witness section where applicable
- Format as professional legal document

CRITICAL - FORMATTING RULES:
- DO NOT use LaTeX formatting under any circumstances
- NO LaTeX commands like \\textbf, \\section, \\begin, \\end, $, etc.
- Use PLAIN TEXT only
- Use ALL CAPS for main headings
- Use simple numbering (1. 2. 3. or 1.1, 1.2, etc.)
- Use dashes (---) for separators
- Use asterisks (*) or dashes (-) for bullet points
- Keep it clean and readable as plain text"""

    user_prompt = f"""Generate a comprehensive {contract_info['name']} with the following details:

CONTRACT TYPE: {contract_info['name']}
APPLICABLE LAW: {contract_info['law_reference']}
LANGUAGE: {language}

PROVIDED INFORMATION:
{form_data_text}

DOCUMENT REQUIREMENTS:
1. Begin with proper title and document date
2. Include recitals/whereas clauses
3. Define all key terms in a Definitions section
4. Include all standard clauses for this contract type
5. Add Indian-specific compliance clauses:
   - Stamp duty clause
   - Jurisdiction clause (courts in the specified state)
   - Governing law (Laws of India)
   - Dispute resolution (Arbitration under Arbitration and Conciliation Act, 1996)
6. Include confidentiality and non-disclosure provisions where applicable
7. Add indemnification clauses
8. Include termination provisions
9. Add force majeure clause
10. Include entire agreement clause
11. Add severability clause
12. Include amendment provisions
13. Add signature blocks with space for:
    - Signature
    - Name
    - Designation
    - Date
    - Place
14. Include witness section

Generate a complete, professional contract ready for legal review.
Format the output in clean, well-structured PLAIN TEXT with proper headings and numbering.

IMPORTANT: Do NOT use LaTeX, mathematical notation, or any special formatting symbols. Output must be plain text only."""

    return system_prompt, user_prompt


async def generate_contract_with_ai(contract_type: str, form_data: Dict[str, Any], language: str = "English") -> str:
    """Generate contract content using OpenAI"""

    system_prompt, user_prompt = build_contract_prompt(contract_type, form_data, language)

    try:
        client = get_client()
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,  # Lower temperature for consistent legal language
            max_tokens=8000,
        )

        return response.choices[0].message.content
    except ValueError as e:
        # Configuration error (API key not set)
        raise Exception(str(e))
    except Exception as e:
        error_msg = str(e)
        # Provide more helpful error messages
        if "invalid_api_key" in error_msg.lower() or "incorrect api key" in error_msg.lower():
            raise Exception("Invalid OpenAI API key. Please check your configuration.")
        if "rate_limit" in error_msg.lower():
            raise Exception("OpenAI rate limit exceeded. Please try again in a moment.")
        if "insufficient_quota" in error_msg.lower():
            raise Exception("OpenAI API quota exceeded. Please check your billing.")
        raise Exception(f"Failed to generate contract: {error_msg}")


# ============== PDF Generation ==============

def generate_pdf(content: str, title: str) -> bytes:
    """Generate PDF from contract content using reportlab"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF generation requires reportlab. Install with: pip install reportlab"
        )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=30
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceBefore=6,
        spaceAfter=6,
        leading=14
    )

    story = []

    # Add title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.25 * inch))

    # Process content - split by lines and handle formatting
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue

        # Check for headings (lines starting with numbers or in caps)
        if line.startswith('#'):
            # Markdown heading
            line = line.lstrip('#').strip()
            story.append(Paragraph(line, heading_style))
        elif line.isupper() or (len(line) < 100 and line.endswith(':')):
            story.append(Paragraph(line, heading_style))
        else:
            # Escape special characters for reportlab
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(line, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ============== DOCX Generation ==============

def generate_docx(content: str, title: str) -> bytes:
    """Generate DOCX from contract content using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="DOCX generation requires python-docx. Install with: pip install python-docx"
        )

    doc = Document()

    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Check for headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.isupper() or (len(line) < 100 and line.endswith(':')):
            heading = doc.add_heading(line, level=2)
        else:
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============== API Endpoints ==============

@router.get("/types", response_model=ContractTypesListResponse)
async def list_contract_types():
    """
    List all available contract types grouped by category.
    Returns contract templates with metadata for Indian law compliance.
    """
    categories = get_contracts_by_category()
    return ContractTypesListResponse(
        categories=categories,
        total_count=len(CONTRACT_TYPES)
    )


@router.get("/types/{type_id}", response_model=ContractTypeResponse)
async def get_contract_type_details(type_id: str):
    """
    Get detailed information about a specific contract type.
    Returns required fields, description, and Indian law references.
    """
    if type_id not in CONTRACT_TYPES:
        raise HTTPException(
            status_code=404,
            detail=f"Contract type '{type_id}' not found. Available types: {', '.join(CONTRACT_TYPES.keys())}"
        )

    contract = CONTRACT_TYPES[type_id]
    return ContractTypeResponse(
        id=type_id,
        name=contract["name"],
        category=contract["category"],
        description=contract["description"],
        law_reference=contract["law_reference"],
        fields=contract["fields"],
        field_labels=contract.get("field_labels", {}),
        field_types=contract.get("field_types", {})
    )


@router.post("/generate", response_model=GeneratedContractResponse)
async def generate_contract(request: ContractGenerateRequest):
    """
    Generate a contract using AI based on the contract type and provided form data.

    The AI will generate a comprehensive, legally compliant contract based on:
    - Contract type template
    - User-provided field data
    - Indian law requirements
    - Standard legal clauses
    """
    # Normalize contract type: convert hyphens to underscores for frontend compatibility
    normalized_type = request.contract_type.replace("-", "_")

    # Also handle some common mappings for frontend contract types
    contract_type_mapping = {
        "founders_agreement": "founders_agreement",
        "shareholder_agreement": "founders_agreement",  # Map to closest available
        "esop_scheme": "founders_agreement",  # Map to closest available
        "employment_contract": "employment",
        "consultant_agreement": "consultant_agreement",
        "internship_agreement": "employment",  # Map to closest available
        "offer_letter": "employment",  # Map to closest available
        "mou_loi": "service_agreement",  # Map to closest available
        "vendor_agreement": "service_agreement",  # Map to closest available
        "partnership_deed": "service_agreement",  # Map to closest available
        "website_disclaimer": "terms_of_service",  # Map to closest available
    }

    # Try normalized type first, then mapping, then original
    contract_type = normalized_type
    if contract_type not in CONTRACT_TYPES:
        contract_type = contract_type_mapping.get(normalized_type, normalized_type)

    # Validate contract type
    if contract_type not in CONTRACT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid contract type '{request.contract_type}'. Available types: {', '.join(CONTRACT_TYPES.keys())}"
        )

    contract_info = CONTRACT_TYPES[contract_type]

    # Flexible field validation: check if form_data has any meaningful content
    # The frontend uses different field names (camelCase) than backend (snake_case)
    # So we just ensure there's at least some data provided
    if not request.form_data or len(request.form_data) == 0:
        raise HTTPException(
            status_code=400,
            detail="No form data provided. Please fill in the required fields."
        )

    # Check if there's at least one non-empty value
    has_content = any(
        value is not None and value != "" and value is not False
        for value in request.form_data.values()
    )

    if not has_content:
        raise HTTPException(
            status_code=400,
            detail="Form data is empty. Please fill in at least some fields."
        )

    # Generate contract with AI
    try:
        content = await generate_contract_with_ai(
            contract_type=contract_type,
            form_data=request.form_data,
            language=request.language
        )
    except Exception as e:
        import traceback
        error_detail = str(e)
        # Check for specific OpenAI errors
        if "api_key" in error_detail.lower() or "authentication" in error_detail.lower():
            raise HTTPException(
                status_code=500,
                detail="OpenAI API key is not configured properly. Please check the server configuration."
            )
        if "rate_limit" in error_detail.lower():
            raise HTTPException(
                status_code=429,
                detail="OpenAI API rate limit exceeded. Please try again in a few moments."
            )
        print(f"Contract generation error: {error_detail}")
        print(traceback.format_exc())
        raise HTTPException(
            status_code=500,
            detail=f"Contract generation failed: {error_detail}"
        )

    # Create contract record
    contract_id = str(uuid4())
    created_at = datetime.now(timezone.utc).isoformat()

    contract_record = {
        "id": contract_id,
        "contract_type": contract_type,
        "contract_name": contract_info["name"],
        "content": content,
        "form_data": request.form_data,
        "created_at": created_at,
        "law_reference": contract_info["law_reference"],
        "status": "generated",
        "language": request.language
    }

    # Store in memory (replace with Supabase in production)
    generated_contracts_store[contract_id] = contract_record

    # Try to store in Supabase if available
    if supabase:
        try:
            supabase.table("generated_contracts").insert({
                "id": contract_id,
                "contract_type": contract_type,
                "contract_name": contract_info["name"],
                "content": content,
                "form_data": request.form_data,
                "law_reference": contract_info["law_reference"],
                "status": "generated",
                "language": request.language
            }).execute()
        except Exception as e:
            # Log error but don't fail - contract is still in memory
            print(f"Warning: Failed to save contract to Supabase: {e}")

    return GeneratedContractResponse(**contract_record)


@router.get("/{contract_id}", response_model=GeneratedContractResponse)
async def get_generated_contract(contract_id: str):
    """
    Get a previously generated contract by its ID.
    """
    # Try memory store first
    if contract_id in generated_contracts_store:
        return GeneratedContractResponse(**generated_contracts_store[contract_id])

    # Try Supabase
    if supabase:
        try:
            response = supabase.table("generated_contracts").select("*").eq("id", contract_id).execute()
            if response.data:
                contract = response.data[0]
                return GeneratedContractResponse(
                    id=contract["id"],
                    contract_type=contract["contract_type"],
                    contract_name=contract["contract_name"],
                    content=contract["content"],
                    form_data=contract["form_data"],
                    created_at=contract["created_at"],
                    law_reference=contract["law_reference"],
                    status=contract.get("status", "generated")
                )
        except Exception as e:
            print(f"Warning: Failed to fetch contract from Supabase: {e}")

    raise HTTPException(status_code=404, detail="Contract not found")


@router.get("/{contract_id}/download")
async def download_contract(
    contract_id: str,
    format: str = Query(default="pdf", description="Download format: pdf or docx")
):
    """
    Download a generated contract as PDF or DOCX.

    Query Parameters:
    - format: 'pdf' or 'docx' (default: pdf)
    """
    # Validate format
    format = format.lower()
    if format not in ["pdf", "docx"]:
        raise HTTPException(
            status_code=400,
            detail="Invalid format. Supported formats: pdf, docx"
        )

    # Get contract
    contract = None

    # Try memory store first
    if contract_id in generated_contracts_store:
        contract = generated_contracts_store[contract_id]

    # Try Supabase
    if not contract and supabase:
        try:
            response = supabase.table("generated_contracts").select("*").eq("id", contract_id).execute()
            if response.data:
                contract = response.data[0]
        except Exception as e:
            print(f"Warning: Failed to fetch contract from Supabase: {e}")

    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    content = contract["content"]
    title = contract["contract_name"]
    filename = f"{title.replace(' ', '_')}_{contract_id[:8]}"

    try:
        if format == "pdf":
            file_bytes = generate_pdf(content, title)
            media_type = "application/pdf"
            filename = f"{filename}.pdf"
        else:
            file_bytes = generate_docx(content, title)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{filename}.docx"

        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate {format.upper()} file: {str(e)}"
        )


@router.get("/", response_model=List[ContractListItem])
async def list_generated_contracts(
    limit: int = Query(default=20, le=100, description="Number of contracts to return"),
    offset: int = Query(default=0, ge=0, description="Offset for pagination")
):
    """
    List all generated contracts with pagination.
    """
    contracts = []

    # Get from memory store
    for contract_id, contract in list(generated_contracts_store.items())[offset:offset + limit]:
        contracts.append(ContractListItem(
            id=contract["id"],
            contract_type=contract["contract_type"],
            contract_name=contract["contract_name"],
            created_at=contract["created_at"],
            status=contract.get("status", "generated")
        ))

    # If we need more and Supabase is available
    if len(contracts) < limit and supabase:
        try:
            response = (
                supabase.table("generated_contracts")
                .select("id, contract_type, contract_name, created_at, status")
                .order("created_at", desc=True)
                .range(offset, offset + limit - 1)
                .execute()
            )

            # Add contracts not already in memory
            memory_ids = set(generated_contracts_store.keys())
            for contract in response.data:
                if contract["id"] not in memory_ids:
                    contracts.append(ContractListItem(
                        id=contract["id"],
                        contract_type=contract["contract_type"],
                        contract_name=contract["contract_name"],
                        created_at=contract["created_at"],
                        status=contract.get("status", "generated")
                    ))
        except Exception as e:
            print(f"Warning: Failed to fetch contracts from Supabase: {e}")

    return contracts[:limit]
