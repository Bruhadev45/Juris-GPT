from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import io
from app.database import supabase
from app.config import settings
from uuid import UUID
from typing import Optional


def markdown_to_docx(markdown_content: str, filename: str) -> bytes:
    """Convert markdown content to Word document (.docx)"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse markdown and convert to docx
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith(tuple(str(i) + '. ' for i in range(1, 100))):
            # Numbered list
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


async def upload_document_to_storage(file_content: bytes, filename: str, folder: str = "documents") -> str:
    """Upload document to Supabase Storage"""
    try:
        file_path = f"{folder}/{filename}"
        
        response = supabase.storage.from_("documents").upload(
            file_path,
            file_content,
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )
        
        # Get public URL
        url_response = supabase.storage.from_("documents").get_public_url(file_path)
        return url_response
    except Exception as e:
        raise Exception(f"Failed to upload document: {str(e)}")


async def create_document_version(
    matter_id: UUID,
    content: str,
    version: int = 1,
    is_final: bool = False
) -> dict:
    """Create a new document version in database"""
    # Convert markdown to docx
    filename = f"founder_agreement_{matter_id}_v{version}.docx"
    docx_content = markdown_to_docx(content, filename)
    
    # Upload to storage
    storage_url = await upload_document_to_storage(docx_content, filename)
    
    # Save to database
    document_data = {
        "matter_id": str(matter_id),
        "content": content,
        "version": version,
        "is_final": is_final,
        "storage_url": storage_url,
        "file_name": filename,
    }
    
    response = supabase.table("documents").insert(document_data).execute()
    
    if not response.data:
        raise Exception("Failed to create document record")
    
    return response.data[0]


async def get_latest_document(matter_id: UUID) -> Optional[dict]:
    """Get the latest document version for a matter"""
    response = supabase.table("documents").select("*").eq("matter_id", str(matter_id)).order("version", desc=True).limit(1).execute()
    
    if response.data:
        return response.data[0]
    return None
